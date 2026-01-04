import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

def generate_pdf_report(results):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph(f"API Comparison Report", styles['Title']))
    story.append(Spacer(1, 12))

    # Summary Info
    story.append(Paragraph(f"Timestamp: {results['timestamp']}", styles['Normal']))
    story.append(Paragraph(f"Environments: {', '.join(results['envs'])}", styles['Normal']))
    story.append(Paragraph(f"Total APIs: {results['api_count']}", styles['Normal']))
    
    pass_rate = (results['consistent_count'] / results['api_count'] * 100) if results['api_count'] > 0 else 0
    story.append(Paragraph(f"Overall Pass Rate: {pass_rate:.2f}%", styles['Normal']))
    story.append(Paragraph(f"Consistent: {results['consistent_count']} | Inconsistent: {results['inconsistent_count']} | Error: {results.get('error_count', 0)}", styles['Normal']))
    story.append(Spacer(1, 24))

    # API Details Table
    story.append(Paragraph("Detailed Results", styles['Heading2']))
    story.append(Spacer(1, 12))

    data = [['API Name', 'Status', 'Similarity']]
    for api_id, api_data in results['api_results'].items():
        similarity = api_data.get('similarity', 'N/A')
        if isinstance(similarity, (int, float)):
            similarity = f"{similarity}%"
        data.append([api_data['name'], api_data['overall_status'], similarity])

    t = Table(data, colWidths=[350, 80, 70])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    # Color status cells
    for i in range(1, len(data)):
        status = data[i][1]
        if status == "Consistent":
            t.setStyle(TableStyle([('TEXTCOLOR', (1, i), (1, i), colors.green)]))
        elif status == "Inconsistent":
            t.setStyle(TableStyle([('TEXTCOLOR', (1, i), (1, i), colors.red)]))
        else:
            t.setStyle(TableStyle([('TEXTCOLOR', (1, i), (1, i), colors.orange)]))

    story.append(t)
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_word_report(results):
    document = Document()
    document.add_heading('API Comparison Report', 0)

    document.add_paragraph(f"Timestamp: {results['timestamp']}")
    document.add_paragraph(f"Environments: {', '.join(results['envs'])}")
    document.add_paragraph(f"Total APIs: {results['api_count']}")
    
    pass_rate = (results['consistent_count'] / results['api_count'] * 100) if results['api_count'] > 0 else 0
    document.add_paragraph(f"Overall Pass Rate: {pass_rate:.2f}%")
    document.add_paragraph(f"Consistent: {results['consistent_count']} | Inconsistent: {results['inconsistent_count']} | Error: {results.get('error_count', 0)}")

    document.add_heading('Detailed Results', level=1)

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'API Name'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Similarity'

    for api_id, api_data in results['api_results'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = api_data['name']
        row_cells[1].text = api_data['overall_status']
        similarity = api_data.get('similarity', 'N/A')
        if isinstance(similarity, (int, float)):
            similarity = f"{similarity}%"
        row_cells[2].text = similarity

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
